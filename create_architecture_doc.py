#!/usr/bin/env python3
"""
Generate Architecture Diagrams Document in .docx format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_color(doc, text, level, color_rgb=None):
    """Add a heading with optional custom color"""
    heading = doc.add_heading(text, level=level)
    if color_rgb:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color_rgb)
    return heading

def add_colored_paragraph(doc, text, color_rgb=None, bold=False, size=None):
    """Add a paragraph with optional color and formatting"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
        if bold:
            run.font.bold = True
        if size:
            run.font.size = Pt(size)
    return para

def create_architecture_doc():
    """Create the architecture document"""
    
    doc = Document()
    doc.add_heading('QShield AI - System Architecture', 0)
    doc.add_paragraph('Comprehensive Architecture Diagrams and Documentation')
    doc.add_paragraph('v1.0 | March 18, 2026')
    
    # Table of Contents
    doc.add_page_break()
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Chrome Extension Architecture',
        '2. Frontend (React + Vite) Architecture',
        '3. Full System Architecture',
        '4. Data Flow & Integration',
        '5. Component Specifications',
        '6. Deployment Architecture'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    # ============= SECTION 1: Chrome Extension Architecture =============
    doc.add_page_break()
    doc.add_heading('1. Chrome Extension Architecture', 1)
    
    doc.add_heading('Overview', 2)
    doc.add_paragraph(
        'The Chrome Extension (Manifest V3) operates as a browser-based phishing detector that runs '
        'alongside the user\'s browsing experience. It extracts links from web pages and Gmail, '
        'analyzes them through the backend API, and displays risk assessments.'
    )
    
    doc.add_heading('Architecture Diagram', 2)
    doc.add_paragraph(
        '┌─────────────────────────────────────────────────────────────┐\n'
        '│                    CHROME EXTENSION (Manifest V3)           │\n'
        '├─────────────────────────────────────────────────────────────┤\n'
        '│                                                             │\n'
        '│  ┌──────────────────────────────────────────────────────┐  │\n'
        '│  │                  UI LAYER                            │  │\n'
        '│  ├──────────────────────────────────────────────────────┤  │\n'
        '│  │  • popup.html/js      → Analysis UI                 │  │\n'
        '│  │  • options.html/js    → Settings Management         │  │\n'
        '│  │  • popup.css          → Styling                     │  │\n'
        '│  └──────────────────────────────────────────────────────┘  │\n'
        '│                           ↓                                 │\n'
        '│  ┌──────────────────────────────────────────────────────┐  │\n'
        '│  │              CONTENT SCRIPTS                         │  │\n'
        '│  ├──────────────────────────────────────────────────────┤  │\n'
        '│  │  • content-script.js     → Extract links from pages │  │\n'
        '│  │  • gmail-integration.js  → Scan Gmail inbox         │  │\n'
        '│  │  • inject.js             → Page injection           │  │\n'
        '│  └──────────────────────────────────────────────────────┘  │\n'
        '│                           ↓                                 │\n'
        '│  ┌──────────────────────────────────────────────────────┐  │\n'
        '│  │        BACKGROUND SERVICE WORKER                    │  │\n'
        '│  ├──────────────────────────────────────────────────────┤  │\n'
        '│  │  • background.js         → Message routing          │  │\n'
        '│  │  • chrome.storage.sync   → Config storage           │  │\n'
        '│  │  • API communication     → Backend calls            │  │\n'
        '│  └──────────────────────────────────────────────────────┘  │\n'
        '│                           ↓                                 │\n'
        '│              API ENDPOINTS (Backend)                        │\n'
        '│  • POST /analyze          → Main analysis              │\n'
        '│  • GET /health            → Health check               │\n'
        '│  • POST /report-phishing  → Report endpoints           │\n'
        '└─────────────────────────────────────────────────────────────┘',
        style='Normal'
    )
    
    doc.add_heading('Key Components', 2)
    
    components = [
        ('popup.html/js', 'User-facing UI for analyzing pages. Displays risk scores, detected links, and threat details.'),
        ('background.js', 'Service worker handling message routing between content scripts and backend API.'),
        ('content-script.js', 'Injects into all web pages to extract links and send them for analysis.'),
        ('gmail-integration.js', 'Specialized script for Gmail inbox scanning and email security checks.'),
        ('options.html/js', 'Settings page allowing users to configure backend URL and preferences.'),
        ('chrome.storage.sync', 'Cloud storage for persisting user settings across devices.'),
    ]
    
    for comp_name, comp_desc in components:
        p = doc.add_paragraph(style='List Bullet')
        p_format = p.paragraph_format
        run = p.add_run(f'{comp_name}: ')
        run.bold = True
        p.add_run(comp_desc)
    
    doc.add_heading('Data Flow', 3)
    flow_steps = [
        'User clicks extension icon → Popup opens',
        'User clicks "Analyze" button → popup.js triggered',
        'Content script extracts all links from current page',
        'Message sent to background.js with link data',
        'background.js validates data and calls backend /analyze',
        'Backend returns risk scores and threat analysis',
        'Results displayed in popup UI with visual indicators',
        'User can view detailed analysis, report phishing, or export'
    ]
    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')
    
    # ============= SECTION 2: Frontend Architecture =============
    doc.add_page_break()
    doc.add_heading('2. Frontend (React + Vite) Architecture', 1)
    
    doc.add_heading('Overview', 2)
    doc.add_paragraph(
        'The frontend is a modern React application built with Vite for fast development and production builds. '
        'It provides comprehensive phishing analysis visualization with 16 interactive components.'
    )
    
    doc.add_heading('Architecture Diagram', 2)
    doc.add_paragraph(
        '┌──────────────────────────────────────────────────────────────┐\n'
        '│              FRONTEND (React + Vite)                         │\n'
        '├──────────────────────────────────────────────────────────────┤\n'
        '│                                                              │\n'
        '│  ┌────────────────────────────────────────────────────────┐ │\n'
        '│  │              App.jsx (Main Component)                 │ │\n'
        '│  ├────────────────────────────────────────────────────────┤ │\n'
        '│  │  State Management:                                    │ │\n'
        '│  │  • message, messageType, loading, results, error     │ │\n'
        '│  │  • timeline, voiceTranscript, threatLocation         │ │\n'
        '│  └────────────────────────────────────────────────────────┘ │\n'
        '│                          ↓                                   │\n'
        '│  ┌────────────────────────────────────────────────────────┐ │\n'
        '│  │           INPUT COMPONENTS (User Input)              │ │\n'
        '│  ├────────────────────────────────────────────────────────┤ │\n'
        '│  │  • InputBox.jsx       → Text message input           │ │\n'
        '│  │  • VoiceInput.jsx     → Speech-to-text via Web API  │ │\n'
        '│  │  • AnalyzeButton.jsx  → Trigger analysis            │ │\n'
        '│  └────────────────────────────────────────────────────────┘ │\n'
        '│                          ↓                                   │\n'
        '│  ┌────────────────────────────────────────────────────────┐ │\n'
        '│  │      ANALYSIS & VISUALIZATION (16 Components)        │ │\n'
        '│  ├────────────────────────────────────────────────────────┤ │\n'
        '│  │  Display:              Utility:                       │ │\n'
        '│  │  • RiskGauge.jsx       • ErrorBoundary.jsx            │ │\n'
        '│  │  • RadarChart          • LoadingSpinner.jsx           │ │\n'
        '│  │  • ThreatTimeline.jsx  • PDF Export (jsPDF)          │ │\n'
        '│  │  • ThreatHeatmap.jsx   • Custom Hooks                │ │\n'
        '│  │  • FlagsDisplay.jsx                                   │ │\n'
        '│  │  • TechnicalDetails.jsx                               │ │\n'
        '│  │  • ExplanationCard.jsx                                │ │\n'
        '│  │  • QuantumBadge.jsx                                   │ │\n'
        '│  │  • ManipulationFingerprint.jsx                        │ │\n'
        '│  │  + Additional visualization components               │ │\n'
        '│  └────────────────────────────────────────────────────────┘ │\n'
        '│                          ↓                                   │\n'
        '│              Backend API (localhost:8000)                    │\n'
        '│  • POST /analyze    → Main analysis endpoint            │\n'
        '│  • GET /health      → Health check                     │\n'
        '└──────────────────────────────────────────────────────────────┘',
        style='Normal'
    )
    
    doc.add_heading('Component Structure', 2)
    
    sections_data = {
        'Input Components': [
            'InputBox.jsx - Text message input field',
            'VoiceInput.jsx - Web Speech API integration',
            'AnalyzeButton.jsx - Submit button with loading state'
        ],
        'Display Components': [
            'RiskGauge.jsx - Circular risk score visualization',
            'RadarChart (Recharts) - 6-axis threat breakdown',
            'ThreatTimeline.jsx - Analysis history timeline',
            'ThreatHeatmap.jsx - Geographic threat visualization',
            'FlagsDisplay.jsx - Individual risk flag details',
            'TechnicalDetails.jsx - Deep technical analysis'
        ],
        'Utility Components': [
            'ErrorBoundary.jsx - Error handling wrapper',
            'LoadingSpinner.jsx - Loading state indicator',
            'ExplanationCard.jsx - AI-generated summary',
            'QuantumBadge.jsx - Quantum risk indicator'
        ],
        'Build & Deployment': [
            'Vite - Fast build tool',
            'React 18.2.0 - UI framework',
            'Recharts - Data visualization',
            'jsPDF - PDF generation'
        ]
    }
    
    for section, items in sections_data.items():
        doc.add_heading(section, 3)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Data Flow', 3)
    flow_steps = [
        'User enters text or speaks → InputBox/VoiceInput updates state',
        'Click "Analyze" → AnalyzeButton validates and triggers fetch',
        'POST request to /analyze endpoint with message and type',
        'Backend processes and returns JSON with 16 risk scores',
        'App.jsx updates state with results',
        'Components receive props and render visualizations',
        'User can view: risk gauge, radar chart, flags, timeline, heatmap',
        'Export to PDF → jsPDF generates downloadable report'
    ]
    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')
    
    # ============= SECTION 3: Full System Architecture =============
    doc.add_page_break()
    doc.add_heading('3. Full System Architecture', 1)
    
    doc.add_heading('System Overview', 2)
    doc.add_paragraph(
        '┌──────────────────────────────────────────────────────────────────┐\n'
        '│                        CLIENT LAYER                             │\n'
        '├──────────────────────────────────────────────────────────────────┤\n'
        '│                                                                  │\n'
        '│  Chrome Extension      Website (React)      Android App         │\n'
        '│  (Manifest V3)         (Vite + Tailwind)    (Kotlin)            │\n'
        '│                                                                  │\n'
        '│  Popup UI              React Components     Jetpack Compose     │\n'
        '│  │                     │                    │                    │\n'
        '│  content-script.js     Voice Input          MVVM ViewModel       │\n'
        '│  │                     │                    │                    │\n'
        '│  background.js         State Management     Room Database        │\n'
        '│  │                     │                    │                    │\n'
        '│  chrome.storage        Error Handling       Retrofit Client      │\n'
        '│  └─────────────────────┴────────────────────┴────────────────────┘\n'
        '│                          HTTP/REST API\n'
        '├──────────────────────────────────────────────────────────────────┤\n'
        '│                    BACKEND SERVICES (FastAPI)                    │\n'
        '├──────────────────────────────────────────────────────────────────┤\n'
        '│                                                                  │\n'
        '│  POST /analyze → Main Analysis Engine                           │\n'
        '│  GET /health   → Health Check                                  │\n'
        '│  POST /report-phishing → Report Endpoints                      │\n'
        '│                  ↓                                               │\n'
        '│  ┌──────────────────────────────────────────────────────────┐   │\n'
        '│  │         ANALYSIS ENGINE (16 Features)                   │   │\n'
        '│  ├──────────────────────────────────────────────────────────┤   │\n'
        '│  │  • URL Analysis              • Quantum Risk Assessment   │   │\n'
        '│  │  • Domain Verification       • Behavior Analysis        │   │\n'
        '│  │  • SSL/TLS Checking          • Threat Intelligence     │   │\n'
        '│  │  • Header Analysis           • Google Generative AI    │   │\n'
        '│  │  • Content Inspection        • 8 additional features   │   │\n'
        '│  └──────────────────────────────────────────────────────────┘   │\n'
        '│                          ↓                                       │\n'
        '│  ┌──────────────────────────────────────────────────────────┐   │\n'
        '│  │              DATA STORAGE                               │   │\n'
        '│  ├──────────────────────────────────────────────────────────┤   │\n'
        '│  │  • SQLite Backend    → Analysis history                │   │\n'
        '│  │  • Room DB (Mobile)  → Local offline storage           │   │\n'
        '│  │  • Chrome Storage    → Extension settings              │   │\n'
        '│  └──────────────────────────────────────────────────────────┘   │\n'
        '└──────────────────────────────────────────────────────────────────┘',
        style='Normal'
    )
    
    doc.add_heading('Integration Points', 2)
    
    # Create table for integration points
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Endpoint'
    hdr_cells[2].text = 'Data Type'
    
    # Data rows
    integrations = [
        ('Chrome Extension', 'POST /analyze', 'URL-encoded or JSON'),
        ('React Frontend', 'POST /analyze', 'JSON Payload'),
        ('Android App', 'POST /analyze', 'JSON via Retrofit'),
        ('All Clients', 'GET /health + POST /report', 'Health check & Reports')
    ]
    
    for i, (component, endpoint, data_type) in enumerate(integrations, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = component
        row_cells[1].text = endpoint
        row_cells[2].text = data_type
    
    # ============= SECTION 4: Data Flow =============
    doc.add_page_break()
    doc.add_heading('4. Data Flow & Integration', 1)
    
    doc.add_heading('Request Flow', 2)
    doc.add_paragraph(
        '1. Client sends POST request to /analyze\n'
        '   {\n'
        '     "message": "Suspicious email content or URL",\n'
        '     "message_type": "email|website|sms"\n'
        '   }\n\n'
        '2. Backend receives and validates input\n\n'
        '3. Core engine runs 16 detection features\n\n'
        '4. Results aggregated into JSON response\n'
        '   {\n'
        '     "final_risk_score": 85.5,\n'
        '     "risk_level": "HIGH",\n'
        '     "all_flags": [...],\n'
        '     "feature_scores": {...},\n'
        '     "explanation": "AI summary",\n'
        '     "timestamp": "2026-03-18T10:30:00"\n'
        '   }\n\n'
        '5. Client receives and stores locally\n\n'
        '6. UI components render analysis visualizations'
    )
    
    doc.add_heading('Response Format', 2)
    doc.add_paragraph(
        'Key Response Fields:\n'
        '• final_risk_score: 0-100 numerical score\n'
        '• risk_level: "SAFE" | "LOW" | "MEDIUM" | "HIGH" | "CRITICAL"\n'
        '• all_flags: Array of detected threats\n'
        '• feature_scores: Breakdown of all 16 features\n'
        '• explanation: AI-generated summary\n'
        '• quantum_score: Quantum risk assessment (0-100)\n'
        '• suggested_action: Recommended user action'
    )
    
    # ============= SECTION 5: Component Specifications =============
    doc.add_page_break()
    doc.add_heading('5. Component Specifications', 1)
    
    components_spec = [
        {
            'name': 'Chrome Extension - popup.js',
            'responsibility': 'Handles user interaction in extension popup',
            'methods': ['checkServerStatus()', 'analyzeLink()', 'displayResults()', 'exportReport()'],
            'dependencies': ['backend.js', 'chrome.storage', 'Fetch API']
        },
        {
            'name': 'Chrome Extension - background.js',
            'responsibility': 'Service worker for message routing and API calls',
            'methods': ['analyzeLink()', 'reportPhishing()', 'getAnalysis()', 'onMessage()'],
            'dependencies': ['chrome.runtime', 'Fetch API', 'chrome.storage']
        },
        {
            'name': 'Frontend - App.jsx',
            'responsibility': 'Main component managing all state and data flow',
            'methods': ['handleAnalyze()', 'startVoiceInput()', 'exportToPDF()', 'handleError()'],
            'dependencies': ['React Hooks', 'Fetch API', 'Recharts', 'jsPDF']
        },
        {
            'name': 'Frontend - RiskGauge.jsx',
            'responsibility': 'Visual display of risk score in gauge format',
            'methods': ['calculatePercentage()', 'getRiskColor()', 'render()'],
            'dependencies': ['React', 'Recharts']
        },
        {
            'name': 'Android - MainActivity.kt',
            'responsibility': 'App entry point, initialization of components',
            'methods': ['onCreate()', 'setupRetrofit()', 'initializeDatabase()'],
            'dependencies': ['Retrofit', 'Room', 'Jetpack Compose']
        }
    ]
    
    for comp in components_spec:
        doc.add_heading(comp['name'], 3)
        doc.add_paragraph(f"Responsibility: {comp['responsibility']}")
        
        doc.add_heading('Key Methods', 4)
        for method in comp['methods']:
            doc.add_paragraph(method, style='List Bullet')
        
        doc.add_heading('Dependencies', 4)
        for dep in comp['dependencies']:
            doc.add_paragraph(dep, style='List Bullet')
    
    # ============= SECTION 6: Deployment Architecture =============
    doc.add_page_break()
    doc.add_heading('6. Deployment Architecture', 1)
    
    doc.add_heading('Development Environment', 2)
    doc.add_paragraph(
        'Local Development:\n'
        '• Backend: FastAPI running on localhost:8000\n'
        '• Frontend: React dev server running on localhost:5173/5174\n'
        '• Extension: Loaded unpacked in Chrome dev mode\n'
        '• Android: Connected to emulator or physical device\n'
        '• Database: SQLite for backend storage'
    )
    
    doc.add_heading('Production Deployment', 2)
    doc.add_paragraph(
        'Vercel Deployment:\n'
        '• Frontend (React): Deployed to Vercel CDN\n'
        '  - URL: https://squid-game-073.vercel.app\n'
        '  - Built with Vite\n'
        '  - Static files served globally\n\n'
        '• Backend (FastAPI): Serverless functions on Vercel\n'
        '  - Using Mangum ASGI adapter\n'
        '  - Python 3.11 runtime\n'
        '  - API routes: /api/*\n'
        '  - URL: https://squid-game-073.vercel.app/api\n\n'
        '• Chrome Extension:\n'
        '  - Published on Chrome Web Store\n'
        '  - Auto-updates via Chrome extension system\n'
        '  - Configured to point to Vercel backend\n\n'
        '• Android App:\n'
        '  - Built as APK for installation\n'
        '  - Configured with Vercel backend URL\n'
        '  - Can be distributed via Google Play Store'
    )
    
    doc.add_heading('Environment Variables', 2)
    env_vars = {
        'BACKEND_URL': 'https://squid-game-073.vercel.app/api',
        'FRONTEND_URL': 'https://squid-game-073.vercel.app',
        'API_TIMEOUT': '30 seconds',
        'MAX_RETRIES': '3'
    }
    
    for key, value in env_vars.items():
        doc.add_paragraph(f'{key} = {value}', style='List Bullet')
    
    # Final page with summary
    doc.add_page_break()
    doc.add_heading('Summary', 1)
    
    doc.add_paragraph(
        'QShield AI implements a distributed phishing detection system with three '
        'independent clients (Chrome Extension, React Website, Android App) all connecting '
        'to a unified FastAPI backend. Each client can operate independently while sharing '
        'analysis results and threat intelligence through the centralized API.'
    )
    
    doc.add_heading('Key Advantages', 2)
    advantages = [
        'Multi-platform coverage (Web, Extension, Mobile)',
        'Centralized threat analysis engine',
        'Real-time risk scoring and visualization',
        'Unified threat intelligence database',
        'Scalable serverless backend on Vercel',
        'Offline support with local caching',
        'User-friendly risk visualizations'
    ]
    for adv in advantages:
        doc.add_paragraph(adv, style='List Bullet')
    
    doc.add_heading('Technology Stack', 2)
    
    tech_table = doc.add_table(rows=5, cols=2)
    tech_table.style = 'Light Grid Accent 1'
    
    tech_headers = tech_table.rows[0].cells
    tech_headers[0].text = 'Layer'
    tech_headers[1].text = 'Technology'
    
    tech_data = [
        ('Frontend', 'React 18.2.0, Vite, Tailwind CSS, Recharts, jsPDF'),
        ('Chrome Extension', 'Manifest V3, JavaScript, Chrome APIs'),
        ('Mobile', 'Kotlin, Jetpack Compose, Room, Retrofit'),
        ('Backend', 'FastAPI 0.104.1, Python, SQLite, Google Generative AI')
    ]
    
    for i, (layer, tech) in enumerate(tech_data, 1):
        cells = tech_table.rows[i].cells
        cells[0].text = layer
        cells[1].text = tech
    
    # Save document
    output_path = r'e:\squid game\qshield-ai\QShield_AI_Architecture.docx'
    doc.save(output_path)
    print(f'✅ Architecture document created: {output_path}')
    return output_path

if __name__ == '__main__':
    try:
        create_architecture_doc()
    except Exception as e:
        print(f'❌ Error: {e}')
        import traceback
        traceback.print_exc()
